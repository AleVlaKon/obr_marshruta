import openpyxl as xl
from docxtpl import DocxTemplate
from docxcompose.composer import Composer

import os
import docx

from start_table import start_table
from WGS_Table import wgs_table
from some_tip import some_tips
from osnov_vid_def import vivodi_v_otchet
from refactor import *


workbook = xl.load_workbook('Ведомость тест.xlsx', data_only=True)
sheet_names = [i for i in workbook.sheetnames if i not in ['Лист1', 'ИД', 'В обсл', 'аб1']]
# sheet_1 = workbook['10']


def context_table(table_cells, sheet):
    """Делает список словарей для таблиц по ключам и номерам столбцов из table_cells для листа sheet"""
    table = []
    for i in range(12, len(sheet['A'])):
        if sheet.cell(row=i, column=1).value not in [None, 'None']:
            table.append({key: sheet.cell(i, table_cells[key]).value for key in table_cells})
    # print(table)
    return table


def change_table_2(table_2: list):
    ''' Редактирует таблицу 2 (замена . на , и добавление хвостовых нулей'''
    # {'km_nach': 45, 'km_kon': 46, 'pokr_i': 47, 'shir_i': 48, 'ball_i': 49,}
    for row in table_2:
        row['km_nach'] = format_float_value(row['km_nach'], 3)
        row['km_kon'] = format_float_value(row['km_kon'], 3)
        row['shir_i'] = format_shirina(row['shir_i'])
        row['ball_i'] = format_float_value(row['ball_i'], 1)


def change_table_3(table_3: list):
    ''' Редактирует таблицу 3 (замена . на , и добавление хвостовых нулей'''
    # {'km': 51, 'ball_i': 52, 'kpr_i': 53, }
    for row in table_3:
        row['ball_i'] = format_float_value(row['ball_i'], 1)
        row['kpr_i'] = format_float_value(row['kpr_i'], 2)


def change_table_4(table_4: list):
    ''' Редактирует таблицу 4 (замена . на , и добавление хвостовых нулей'''
    # {'km': 56, 'kpr_i': 57, 'E_i': 58, }
    for row in table_4:
        row['kpr_i'] = format_float_value(row['kpr_i'], 2)
        if isinstance(row['E_i'], int | float):
            row['E_i'] = f"{row['E_i']:,.0f}"
    

def dob_nuley(cell, nuli):
    if type(cell.value) == int:                      #Добавляет nuli, если cell целое число
        return '{},{}'.format(cell.value, nuli)
    else:
        return str(cell.value).replace('.', ',')


def asphalt(sheet, sheetname, template):
    ''' Создает шаблон для дороги с асфальтобетонным покрытием
        sheet - лист;
        sheetname - имя листа, под которым сохраняется заполненный шаблон
        template - шаблон, который будет заполняться
    '''
    template = DocxTemplate(template)

    
    table_1_cells = {'km': 39, 'defect': 41}
    table_1 = context_table(table_1_cells, sheet)


    table_2_cells = {'km_nach': 45, 'km_kon': 46, 'pokr_i': 47, 'shir_i': 48, 'ball_i': 49,}
    table_2 = context_table(table_2_cells, sheet)
    change_table_2(table_2)

 
    table_3_cells = {'km': 51, 'ball_i': 52, 'kpr_i': 53, }
    table_3 = context_table(table_3_cells, sheet)
    change_table_3(table_3)


    table_4_cells = {'km': 56, 'kpr_i': 57, 'E_i': 58, }
    table_4 = context_table(table_4_cells, sheet)
    change_table_4(table_4)

    
    dor_od = {
        'asf': 'Конструкция дорожной одежды нежесткая, облегченного типа с покрытием из асфальтобетона',
        'sheb': 'Конструкция дорожной одежды нежесткая, переходного типа с щебеночным покрытием',
        'grav': 'Конструкция дорожной одежды нежесткая, переходного типа с гравийным покрытием',
        'shps': 'Конструкция дорожной одежды нежесткая, переходного типа с щебеночно-песчаным покрытием',      
    }
    
    if sheet['B5'].value in ('асфальтобетон', 'асф. бет.', 'асфальтобетонное', 'асф.бет.', 'асф.бет'):
        konstr_do = dor_od['asf']
    elif sheet['B5'].value in ('щебеночное', 'щебень'):
        konstr_do = dor_od['sheb']
    elif sheet['B5'].value in ('гравийное', 'гравий'):
        konstr_do = dor_od['grav']
    elif sheet['B5'].value in ('цементобетон', 'цементобетонное', 'бетонное'):
        konstr_do = dor_od['bet']
    elif sheet['B5'].value in ('ЩПС'):
        konstr_do = dor_od['shps']
    else:
        konstr_do = some_tips(sheet)


    #Declare template variables
    additional_context = {
        'ball_sr': format_float_value(sheet['K2'].value, 1),
        'kpr_sr': format_float_value(sheet['K4'].value, 2),
        'vyvody': sheet['AM2'].value,
        'konstr_do': konstr_do,
        'table_1': table_1,
        'table_2': table_2,
        'table_3': table_3,
        'table_4': table_4,
        'vivod': vivodi_v_otchet(sheet)[0]

        }

    context = return_base_context(sheet)
    context.update(additional_context)
    print(context)
    template.render(context)
    template.save(f'temp/{sheetname}.docx')
    print(f'Маршрут {sheetname} сохранен')


def PGS(sheet, sheetname):
    ''' Создает шаблон для дороги с покрытием из ПГС'''
    template = DocxTemplate('templates/ПГС.docx')

    
    table_1_cells = {'km': 39, 'defect': 41}
    table_1 = context_table(table_1_cells, sheet)

    
    dor_od = {
        'ПГС': 'Конструкция дорожной одежды на всём протяжении – нежесткая низшего типа с песчано-гравийным покрытием',
        'плиты': 'Конструкция дорожной одежды жесткая, капитального типа с покрытием из железобетонных плит',
        'bet': 'Конструкция дорожной одежды жесткая, капитального типа с бетонным покрытием',                       
    }
    #Добавить грунтощебень, скальный грунт, шлак
    if sheet['B5'].value in ('ПГС', ):
        konstr_do = dor_od['ПГС']
    elif sheet['B5'].value in ('плиты',):
        konstr_do = dor_od['плиты']
    elif sheet['B5'].value in ('цементобетон', 'цементобетонное', 'бетонное',):
        konstr_do = dor_od['bet']        


    #Declare template variables
    additional_context = {
        'table_1': table_1,
        'konstr_do': konstr_do,
        }

    context = return_base_context(sheet).update(additional_context)

    template.render(context)
    template.save(f'temp/{sheetname}.docx')
    print(f'Маршрут {sheetname} сохранен')


def Gruntovaya(sheet, sheetname):
    ''' Создает шаблон для грунтовой дороги'''
    template = DocxTemplate('templates/Грунтовая.docx')

    dor_od = {
        'grunt': 'Проезд осуществляется по грунтовой дороге',
        'gr_sheb': 'Проезд осуществляется по дороге с покрытием из грунтощебня',
        'gr_ul': 'Проезд осуществляется по дороге с грунтовым улучшенным покрытием',
    }

    if sheet['B5'].value in ('грунтовая', 'грунт', 'грунтовое'):
        konstr_do = dor_od['grunt']
    elif sheet['B5'].value in ('грунтощебень',):
        konstr_do = dor_od['gr_sheb']
    elif sheet['B5'].value in ('грунтовое улучшенное',):
        konstr_do = dor_od['gr_ul']

    
    #Declare template variables
    additional_context = {'konstr_do': konstr_do}

    context = return_base_context(sheet).update(additional_context)

    template.render(context)
    template.save(f'temp/{sheetname}.docx')
    print(f'Маршрут {sheetname} сохранен')


def zapolnenie(list_cat, zagolovok):
    if list_cat:
        master.add_heading(text=zagolovok, level=2)
        for i in list_cat:
            doc2 = docx.Document(f'temp/{i}.docx')
            composer.append(doc2)
            print(f'Документ{i} добавлен')


print(sheet_names)

for i in sheet_names:
    sheet_name = i
    if workbook[sheet_name]['B5'].value.lower() in ('асфальтобетон', 'асф. бет.', 'асфальтобетонное',
                                            'асф.бет.', 'щебеночное', 'щебень', 'гравийное', 'гравий'):
        asphalt(workbook[sheet_name], sheet_name, 'templates/Асфальт.docx')
    elif workbook[sheet_name]['B5'].value.lower() in ('пгс', 'плиты', 'цементобетон', 'цементобетонное', 'бетонное'):
        PGS(workbook[sheet_name], sheet_name)
    elif workbook[sheet_name]['B5'].value.lower() in ('грунтовая', 'грунтощебень', 'грунт', 'грунтовое', 'грунтовое улучшенное'):
        Gruntovaya(workbook[sheet_name], sheet_name)
    else:
        asphalt(workbook[sheet_name], sheet_name, 'templates/Асфальт2.docx')


reg_list = [x for x in sheet_names if workbook[x]['B2'].value == 'региональной']
fed_list = [x for x in sheet_names if workbook[x]['B2'].value == 'федеральной']
mest_list = [x for x in sheet_names if workbook[x]['B2'].value == 'местной']
chastnie_list = [x for x in sheet_names if workbook[x]['B2'].value == 'частной']
lesn_list = [x for x in sheet_names if workbook[x]['B2'].value == 'лесной']
vedom_list = [x for x in sheet_names if workbook[x]['B2'].value == 'ведомственной']

start_table(workbook)

master = docx.Document(f'temp/Шаблон отчета.docx') #Объединение отчетов
composer = Composer(master)
list_of_docs = os.listdir(path='temp')


zapolnenie(fed_list, 'Федеральные автомобильные дороги')
zapolnenie(reg_list, 'Региональные автомобильные дороги')
zapolnenie(mest_list, 'Местные автомобильные дороги')
zapolnenie(chastnie_list, 'Частные автомобильные дороги')
zapolnenie(lesn_list, 'Лесные автомобильные дороги')
zapolnenie(vedom_list, 'Ведомственные автомобильные дороги')

composer.save('Отчет.docx')
wgs_table(workbook)


for file in os.listdir(path='temp'):
    os.remove(f'temp/{file}')
